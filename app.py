# backend.py
import uvicorn
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse, FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
import tempfile, io, os, re, json, base64, hashlib
from typing import List, Tuple, Dict
import fitz  # PyMuPDF
import requests
import pandas as pd
from docx import Document
from io import BytesIO

from sqlalchemy import create_engine, Column, Integer, String, Text, DateTime, Boolean, func 
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
import datetime
from fuzzywuzzy import fuzz
# Track which results have been exported to prevent duplicates
EXPORTED_RESULTS = {}  # key: result_hash -> {"exported": True, "timestamp": datetime}

from urllib.parse import quote_plus




from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker, declarative_base
import os
# Use SQLite instead of MySQL
DATABASE_URL = "sqlite:///./app.db"

engine = create_engine(
    DATABASE_URL, 
    connect_args={"check_same_thread": False}  # Needed for SQLite
)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
from sqlalchemy.orm import declarative_base
Base = declarative_base()


engine = create_engine(DATABASE_URL, pool_pre_ping=True)
SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)
Base = declarative_base()

class Question(Base):
    __tablename__ = "questions"

    id = Column(Integer, primary_key=True, index=True)
    topic = Column(String(255))
    type = Column(String(20))  # MCQ / Descriptive
    question = Column(Text, nullable=False)
    option_a = Column(Text)
    option_b = Column(Text)
    option_c = Column(Text)
    option_d = Column(Text)
    answer = Column(Text)
    descriptive_answer = Column(Text)
    difficulty = Column(String(10))
    pdf_name = Column(String(255))  # <-- new: store PDF filename/original name
    
    created_at = Column(DateTime, default=datetime.datetime.utcnow)
    flagged = Column(Boolean, default=None)  # Change from True to None

# Create table if not exists
Base.metadata.create_all(bind=engine)


# Add this new model for course-specific questions
class CourseQuestion(Base):
    __tablename__ = "course_questions"

    id = Column(Integer, primary_key=True, index=True)
    topic = Column(String(255))
    type = Column(String(20))  # MCQ / Descriptive
    question = Column(Text, nullable=False)
    option_a = Column(Text)
    option_b = Column(Text)
    option_c = Column(Text)
    option_d = Column(Text)
    answer = Column(Text)
    descriptive_answer = Column(Text)
    difficulty = Column(String(10))
    pdf_name = Column(String(255))
    po_id = Column(String(50))
    co_id = Column(String(50))
    program_id = Column(String(100))  # Store selected program
    branch_id = Column(String(100))   # Store selected branch
    academic_year_id = Column(String(100))  # Store selected academic year
    semester_id = Column(String(100))  # Store selected semester
    course_id = Column(String(100))   # Store selected course
    component = Column(String(100))   # Store selected component
    created_at = Column(DateTime, default=datetime.datetime.utcnow)
    flagged = Column(Boolean, default=None)

# Create table if not exists
if not os.path.exists("app.db"):
    Base.metadata.create_all(bind=engine)


import json

def save_questions_to_db(results: dict, pdf_name: str = None):
    db = SessionLocal()
    saved = 0
    skipped = 0
    duplicates = 0

    try:
        for topic, data in (results or {}).items():
            topic_val = topic if topic else None

            # ---------- MCQs ----------
            for mcq in data.get("mcqs", []) if data else []:
                question_text = mcq.get("question") or mcq.get("q")
                if not question_text or not str(question_text).strip():
                    skipped += 1
                    continue

                question_text = str(question_text).strip()

                # 🔴 DUPLICATE CHECK
                exists = db.query(Question).filter(
                    Question.question == question_text,
                    Question.type == "MCQ",
                    Question.pdf_name == pdf_name
                ).first()

                if exists:
                    duplicates += 1
                    continue

                opts = mcq.get("options", []) or []
                q = Question(
                    topic=topic_val,
                    type="MCQ",
                    question=question_text,
                    option_a=opts[0] if len(opts) > 0 else None,
                    option_b=opts[1] if len(opts) > 1 else None,
                    option_c=opts[2] if len(opts) > 2 else None,
                    option_d=opts[3] if len(opts) > 3 else None,
                    answer=mcq.get("answer"),
                    descriptive_answer=None,
                    difficulty=str(mcq.get("difficulty")) if mcq.get("difficulty") else None,
                    pdf_name=pdf_name,
                    created_at=datetime.datetime.utcnow(),
                    flagged=None
                )

                db.add(q)
                saved += 1

            # ---------- DESCRIPTIVE ----------
            for dq in data.get("descriptive", []) if data else []:
                question_text = dq.get("question") or dq.get("q")
                if not question_text or not str(question_text).strip():
                    skipped += 1
                    continue

                question_text = str(question_text).strip()

                # 🔴 DUPLICATE CHECK
                exists = db.query(Question).filter(
                    Question.question == question_text,
                    Question.type == "Descriptive",
                    Question.pdf_name == pdf_name
                ).first()

                if exists:
                    duplicates += 1
                    continue

                q = Question(
                    topic=topic_val,
                    type="Descriptive",
                    question=question_text,
                    option_a=None,
                    option_b=None,
                    option_c=None,
                    option_d=None,
                    answer=None,
                    descriptive_answer=dq.get("answer"),
                    difficulty=str(dq.get("difficulty")) if dq.get("difficulty") else None,
                    pdf_name=pdf_name,
                    created_at=datetime.datetime.utcnow(),
                    flagged=None
                )

                db.add(q)
                saved += 1

        db.commit()

        return {
            "status": "success",
            "saved": saved,
            "duplicates": duplicates,
            "skipped": skipped
        }

    except Exception as e:
        db.rollback()
        return {
            "status": "error",
            "error": str(e)
        }
    finally:
        db.close()


def save_questions_to_course_db(results: dict, pdf_name: str = None, course_data: dict = None):
    """
    Save parsed results into the course_questions table.
    course_data should contain: program_id, branch_id, academic_year_id, 
    semester_id, course_id, component
    """
    db = SessionLocal()
    saved = 0
    skipped = 0

    try:
        print(f"📝 Saving to course_questions table. PDF: {pdf_name}, Course data: {course_data}")
        
        for topic, data in (results or {}).items():
            topic_val = topic if topic is not None else None

            # Save MCQs
            for mcq in data.get("mcqs", []) if data else []:
                question_text = mcq.get("question") or mcq.get("q") or None
                if not question_text or not str(question_text).strip():
                    skipped += 1
                    continue

                opts = mcq.get("options", []) or []
                option_a = opts[0] if len(opts) > 0 else mcq.get("option_a") or None
                option_b = opts[1] if len(opts) > 1 else mcq.get("option_b") or None
                option_c = opts[2] if len(opts) > 2 else mcq.get("option_c") or None
                option_d = opts[3] if len(opts) > 3 else mcq.get("option_d") or None

                answer = mcq.get("answer") or mcq.get("ans") or None
                difficulty = mcq.get("difficulty")
                difficulty = str(difficulty) if difficulty is not None else None

                # Create course question
                q = CourseQuestion(
                    topic=topic_val,
                    type="MCQ",
                    question=str(question_text).strip(),
                    option_a=option_a,
                    option_b=option_b,
                    option_c=option_c,
                    option_d=option_d,
                    answer=answer,
                    descriptive_answer=None,
                    difficulty=difficulty,
                    pdf_name=pdf_name,
                    po_id=None,  # Can be set later
                    co_id=None,  # Can be set later
                    program_id=course_data.get('program_id') if course_data else None,
                    branch_id=course_data.get('branch_id') if course_data else None,
                    academic_year_id=course_data.get('academic_year_id') if course_data else None,
                    semester_id=course_data.get('semester_id') if course_data else None,
                    course_id=course_data.get('course_id') if course_data else None,
                    component=course_data.get('component') if course_data else None,
                    created_at=datetime.datetime.utcnow(),
                    flagged=None
                )
                db.add(q)
                saved += 1
                print(f"✓ Saved MCQ: {question_text[:50]}...")

            # Save Descriptive
            for dq in data.get("descriptive", []) if data else []:
                question_text = dq.get("question") or dq.get("q") or None
                if not question_text or not str(question_text).strip():
                    skipped += 1
                    continue

                descriptive_answer = dq.get("answer") or dq.get("descriptive_answer") or None
                difficulty = dq.get("difficulty")
                difficulty = str(difficulty) if difficulty is not None else None

                # Create course question
                q = CourseQuestion(
                    topic=topic_val,
                    type="Descriptive",
                    question=str(question_text).strip(),
                    option_a=None,
                    option_b=None,
                    option_c=None,
                    option_d=None,
                    answer=None,
                    descriptive_answer=descriptive_answer,
                    difficulty=difficulty,
                    pdf_name=pdf_name,
                    po_id=None,
                    co_id=None,
                    program_id=course_data.get('program_id') if course_data else None,
                    branch_id=course_data.get('branch_id') if course_data else None,
                    academic_year_id=course_data.get('academic_year_id') if course_data else None,
                    semester_id=course_data.get('semester_id') if course_data else None,
                    course_id=course_data.get('course_id') if course_data else None,
                    component=course_data.get('component') if course_data else None,
                    created_at=datetime.datetime.utcnow(),
                    flagged=None
                )
                db.add(q)
                saved += 1
                print(f"✓ Saved Descriptive: {question_text[:50]}...")

        db.commit()
        print(f"✅ Successfully saved {saved} questions to course_questions table")
        return {"status": "success", "saved": saved, "skipped": skipped}

    except Exception as e:
        db.rollback()
        print("❌ DB error in save_questions_to_course_db:", e)
        import traceback
        traceback.print_exc()
        return {"status": "error", "error": str(e)}
    finally:
        db.close()
from dotenv import load_dotenv
load_dotenv()
# OpenRouter Configuration
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")  # Set your API key in environment variable
OPENROUTER_API_URL = "https://openrouter.ai/api/v1/chat/completions"
OPENROUTER_MODEL = "meta-llama/llama-3.3-70b-instruct:free"  # Free model, you can change this

# Headers for OpenRouter API
OPENROUTER_HEADERS = {
    "Authorization": f"Bearer {OPENROUTER_API_KEY}",
    "Content-Type": "application/json",
    "HTTP-Referer": "http://localhost:8000",  # Optional: your site URL
    "X-Title": "MCQ Generator"  # Optional: your app name
}

MODEL = OPENROUTER_MODEL
HOST = "127.0.0.1"
PORT = 8000
# ---------- FASTAPI ----------
app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"], allow_credentials=True)
# Serve static files (put design.html and any assets inside ./static/)
static_dir = os.path.join(os.path.dirname(__file__), "static")
if not os.path.isdir(static_dir):
    os.makedirs(static_dir, exist_ok=True)
app.mount("/static", StaticFiles(directory=static_dir), name="static")

# Serve design.html at root
@app.get("/", response_class=HTMLResponse)
async def index():
    fpath = os.path.join(static_dir, "design.html")
    if os.path.exists(fpath):
        return HTMLResponse(open(fpath, "r", encoding="utf-8").read())
    return HTMLResponse("<h3>Place design.html inside ./static/ and reload.</h3>")
# ---------- IN-MEMORY STATE & STORE ----------
IN_MEMORY_STORE = {}  # key -> {"data": bytes, "name": str, "mime": str}
STATE = {
    "pdf_uploads": 0,
    "last_pdf_hash": None,
    "last_pdf_pages": 0,
    "mcq_count": 0,
    "desc_count": 0
}

def store_result_bytes(key: str, data: bytes, filename: str, mime: str):
    IN_MEMORY_STORE[key] = {"data": data, "name": filename, "mime": mime}

@app.get("/download/{key}")
async def download_key(key: str):
    item = IN_MEMORY_STORE.get(key)
    if not item:
        return JSONResponse({"error": "Not found"}, status_code=404)
    return StreamingResponse(io.BytesIO(item["data"]), media_type=item["mime"],
                             headers={"Content-Disposition": f"attachment; filename={item['name']}"})

@app.get("/status")
async def status():
    """Return counters for the top dashboard (PDF uploads, pages, counts)."""
    return {
        "pdf_uploads": STATE["pdf_uploads"],
        "last_pdf_pages": STATE["last_pdf_pages"],
        "mcq_count": STATE["mcq_count"],
        "desc_count": STATE["desc_count"]
    }

# ---------- UTIL HELPERS (ported from your Streamlit code) ----------
def clean_text(text: str) -> str:
    if text is None:
        return ""
    return re.sub(r"[\x00-\x1F\x7F]", "", str(text))

def detect_index_range(doc, min_section_hits: int = 3, consecutive_break: int = 2) -> Tuple[int, int]:
    scores = []
    has_contents_flags = []
    for pno in range(doc.page_count):
        try:
            text = doc.load_page(pno).get_text("text") or ""
        except Exception:
            text = ""
        low = text.lower()
        has_contents = bool(re.search(r"\btable of contents\b|\bcontents\b", low))
        count_sections = len(re.findall(r"\b\d{1,2}\.\d+\b", text))
        count_leaders = len(re.findall(r"\.{2,}\s*\d+|\s+\d{1,3}\s*$", text, re.M))
        score = count_sections + 0.6 * count_leaders + (5 if has_contents else 0)
        scores.append(score)
        has_contents_flags.append(has_contents)

    if any(has_contents_flags):
        start_idx = next(i for i, f in enumerate(has_contents_flags) if f)
        end_idx = start_idx
        break_count = 0
        for i in range(start_idx + 1, len(scores)):
            if scores[i] >= 1.0:
                end_idx = i
                break_count = 0
            else:
                break_count += 1
                if break_count >= consecutive_break:
                    break
        return (start_idx + 1, end_idx + 1)

    start_idx = None
    for i, s in enumerate(scores):
        if s >= min_section_hits:
            start_idx = i
            break
    if start_idx is None:
        raise ValueError("Could not auto-detect contents/index pages.")

    end_idx = start_idx
    gap = 0
    for i in range(start_idx + 1, len(scores)):
        if scores[i] >= 1.0:
            end_idx = i
            gap = 0
        else:
            gap += 1
            if gap >= consecutive_break:
                break
    return (start_idx + 1, end_idx + 1)
import time, os, requests, json


def call_ollama(prompt: str, model: str = None, timeout: int = 600) -> str:
    try:
        payload = {
            "model": OPENROUTER_MODEL,   # e.g. "meta-llama/llama-3.3-70b-instruct:free"
            "messages": [
                {"role": "user", "content": prompt}
            ]
        }
        resp = requests.post(
            OPENROUTER_API_URL,
            headers=OPENROUTER_HEADERS,
            json=payload,
            timeout=120
        )
        resp.raise_for_status()
        data = resp.json()
        # OpenRouter chat completion shape
        return data["choices"][0]["message"]["content"].strip()
    except Exception as e:
        return f"LOCAL_MODEL_ERROR: {str(e)}"
def generate_mcqs_ollama(topic: str, context: str = "", full_text: str = "", model: str = MODEL, num_qs: int = 0):
    """
    Generate `num_qs` multiple-choice questions with Ollama.
    """
    prompt = f"""
Generate exactly {num_qs} distinct multiple-choice questions for the topic below. For each question include:
- Exactly 4 labeled options A) B) C) D)
- A single-letter correct answer on its own line: Answer: <A/B/C/D>
- (Optional) Difficulty line: Difficulty: <1-5>

Use exactly this format; do not add extra commentary or code fences.

Q1. <question text>
A) <option A>
B) <option B>
C) <option C>
D) <option D>
Answer: <letter>
Difficulty: <1-5>

Q2. <question text>
A) <option A>
B) <option B>
C) <option C>
D) <option D>
Answer: <letter>
Difficulty: <1-5>

Continue this pattern for exactly {num_qs} questions.

Topic: {topic}
Context: {context[:1500]}
"""
    out = call_ollama(prompt, model=model,timeout=600)
    if out.startswith("OLLAMA_ERROR"):
        # If Ollama fails, generate fallback questions to match the requested count
        return generate_fallback_mcqs(topic, num_qs, context)
    
    mcqs = []
    # split by Qn blocks
    blocks = re.split(r'\n(?=Q\d+\.)', out)
    
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines = [ln.rstrip() for ln in block.splitlines() if ln.strip()]
        if not lines:
            continue
        
        q_line_idx = 0
        for idx, ln in enumerate(lines):
            if re.match(r'^\s*Q\d+\.', ln, re.I):
                q_line_idx = idx
                break
            if not re.match(r'^[A-D][\)\.\-:]', ln, re.I) and not re.search(r'(here are|multiple[-\s]?choice|based on the topic)', ln, re.I):
                q_line_idx = idx
                break
        
        q_line = clean_text(lines[q_line_idx])
        opts = []
        opt_end_idx = q_line_idx
        
        for j in range(q_line_idx + 1, len(lines)):
            m = re.match(r'^([A-D])[\)\.\-:]\s*(.*)$', lines[j], re.I)
            if m:
                label = m.group(1).upper()
                text = m.group(2).strip()
                opts.append(f"{label}. {text}")
                opt_end_idx = j
            else:
                # handle continuation lines for previous option (concatenate)
                if opts and lines[j].strip():
                    opts[-1] = opts[-1] + " " + lines[j].strip()
                else:
                    break
        
        answer = ""
        difficulty = ""
        look_start = opt_end_idx + 1
        look_end = min(len(lines), opt_end_idx + 8)
        
        for k in range(look_start, look_end):
            ln = lines[k]
            m_ans = re.search(r'(?i)\b(?:answer|correct)[:\s\-]*\(?\s*([A-D])\s*\)?', ln)
            if m_ans:
                answer = m_ans.group(1).upper()
                continue
            m_diff = re.search(r'(?i)\b(?:difficulty|level)[:\s\-]*\(?\s*([1-5])\s*\)?', ln)
            if m_diff:
                difficulty = m_diff.group(1)
                continue
            m_single = re.match(r'^\s*([A-D])[\)\.\s]*$', ln, re.I)
            if m_single and not answer:
                answer = m_single.group(1).upper()
        
        if not answer:
            m_any = re.search(r'(?i)\banswer[:\s\-]*\(?\s*([A-D])\s*\)?', block)
            if m_any:
                answer = m_any.group(1).upper()
        
        if q_line and len(opts) >= 2:
            mcqs.append({"question": q_line, "options": opts, "answer": answer, "difficulty": difficulty})
    
    # Ensure we have exactly the requested number of questions
    if len(mcqs) < num_qs:
        # Generate additional fallback questions to reach the target count
        additional_needed = num_qs - len(mcqs)
        fallback_mcqs = generate_fallback_mcqs(topic, additional_needed, context)
        mcqs.extend(fallback_mcqs)
    elif len(mcqs) > num_qs:
        # Trim excess questions
        mcqs = mcqs[:num_qs]
    
    return mcqs
def generate_fallback_mcqs(topic: str, num_qs: int, context: str = ""):
    """Generate simple fallback MCQs when Ollama fails or returns insufficient questions."""
    mcqs = []
    for i in range(num_qs):
        mcq = {
            "question": f"What is the main concept of '{topic}'?",
            "options": [
                "A. Fundamental principle discussed in the text",
                "B. Basic terminology introduction", 
                "C. Key application mentioned",
                "D. Core methodology described"
            ],
            "answer": "A",
            "difficulty": "3"
        }
        mcqs.append(mcq)
    return mcqs
def generate_descriptive_with_answers(topic: str, context: str = "", model: str = MODEL, num_qs: int = 0):
    prompt = f"""
Generate exactly {num_qs} descriptive / short-answer / essay-style questions for the topic below.
For each question, also provide:
- Correct answer
- Difficulty level (1-5)

Return exactly in this format:

Q1. <question text>
Answer: <answer text>
Difficulty: <1-5>

Q2. <question text>
Answer: <answer text>
Difficulty: <1-5>

Continue this pattern for exactly {num_qs} questions.

Do not add extra commentary.

Topic: {topic}
Context: {context[:1500]}
"""
    out = call_ollama(prompt, model=model,timeout=600)
    if out.startswith("OLLAMA_ERROR"):
        return generate_fallback_descriptive(topic, num_qs, context)
    
    blocks = re.split(r'\n(?=Q\d+\.)', out)
    results = []
    
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines = block.splitlines()
        question = ""
        answer = ""
        difficulty = ""
        
        for ln in lines:
            ln = ln.strip()
            if ln.lower().startswith("q"):
                question = re.sub(r'^q\d+\.\s*', '', ln, flags=re.I).strip()
            elif ln.lower().startswith("answer:"):
                answer = ln.split(":", 1)[1].strip()
            elif ln.lower().startswith("difficulty:"):
                difficulty = ln.split(":", 1)[1].strip()
        
        if question:
            results.append({"question": question, "answer": answer, "difficulty": difficulty})
    
    # Ensure we have exactly the requested number of questions
    if len(results) < num_qs:
        additional_needed = num_qs - len(results)
        fallback_desc = generate_fallback_descriptive(topic, additional_needed, context)
        results.extend(fallback_desc)
    elif len(results) > num_qs:
        results = results[:num_qs]
    
    return results

def generate_fallback_descriptive(topic: str, num_qs: int, context: str = ""):
    """Generate simple fallback descriptive questions."""
    results = []
    for i in range(num_qs):
        result = {
            "question": f"Explain the key aspects of '{topic}' as discussed in the text.",
            "answer": f"The text discusses various aspects of {topic} including fundamental concepts, applications, and methodologies.",
            "difficulty": "3"
        }
        results.append(result)
    return results

def build_docx_bytes(questions_data: dict) -> bytes:
    doc = Document()
    doc.add_heading("Generated Questions", level=1)
    for topic_title, blocks in questions_data.items():
        doc.add_heading(topic_title, level=2)
        mcqs = blocks.get("mcqs", []) or []
        if mcqs:
            doc.add_paragraph("Multiple Choice Questions:")
            for idx, mcq in enumerate(mcqs, start=1):
                doc.add_paragraph(f"{idx}. {mcq.get('question','')}")
                for opt in mcq.get("options", []):
                    doc.add_paragraph(f"    {opt}")
                ans = mcq.get("answer", "")
                diff = mcq.get("difficulty", "N/A")
                if ans:
                    doc.add_paragraph(f"    Answer: {ans}    Difficulty: {diff}")
                else:
                    doc.add_paragraph(f"    Difficulty: {diff}")
                doc.add_paragraph("")
        descrs = blocks.get("descriptive", []) or []
        if descrs:
            doc.add_paragraph("Descriptive / Short-answer Questions:")
            for idx, dq in enumerate(descrs, start=1):
                if isinstance(dq, dict):
                    q = dq.get("question", "")
                    a = dq.get("answer", "")
                    diff = dq.get("difficulty", "N/A")
                else:
                    q = str(dq)
                    a, diff = "", "N/A"
                doc.add_paragraph(f"{idx}. {q}")
                if a:
                    doc.add_paragraph(f"    Answer: {a}")
                doc.add_paragraph(f"    Difficulty: {diff}")
                doc.add_paragraph("")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def build_dfs_from_questions(questions_data: dict):
    rows = []
    for topic_title, topic_data in questions_data.items():
        for mcq in topic_data.get("mcqs", []):
            opts = mcq.get("options") or []
            rows.append({
                "Topic": topic_title,
                "Type": "MCQ",
                "Question": mcq.get("question", ""),
                "Option A": opts[0] if len(opts) > 0 else "",
                "Option B": opts[1] if len(opts) > 1 else "",
                "Option C": opts[2] if len(opts) > 2 else "",
                "Option D": opts[3] if len(opts) > 3 else "",
                "Answer": mcq.get("answer", ""),
                "Difficulty": mcq.get("difficulty", "N/A"),
                "Descriptive Answer": ""
            })
        for dq in topic_data.get("descriptive", []):
            rows.append({
                "Topic": topic_title,
                "Type": "Descriptive",
                "Question": dq.get("question", ""),
                "Option A": "", "Option B": "", "Option C": "", "Option D": "",
                "Answer": "",
                "Difficulty": dq.get("difficulty", "N/A"),
                "Descriptive Answer": dq.get("answer", "")
            })
    return pd.DataFrame(rows)
# ---------- ENDPOINTS: PDF / TOC / GENERATION ----------
@app.post("/extract_toc")
async def extract_toc(file: UploadFile = File(...)):
    pdf_bytes = await file.read()
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        STATE["last_pdf_pages"] = getattr(doc, "page_count", 0)
        doc.close()
        
        # Extract topics from PDF
        topics = extract_topics_from_pdf(pdf_bytes)
        
        # Format for frontend
        matches = []
        for idx, topic in enumerate(topics):
            matches.append({
                "id": f"topic_{idx}",  # Unique ID
                "title": topic["title"],
                "page": topic["page"],
                "font_size": topic.get("font_size", 12.0),
                "confidence": topic.get("confidence", 0.5),
                "type": topic.get("type", "heading")
            })
        
        return {
            "status": "success", 
            "matches": matches,
            "topics_count": len(matches),
            "pages": STATE["last_pdf_pages"]
        }
    except Exception as e:
        return {"status": "error", "error": str(e)}

@app.post("/generate_pdf_mcqs")
async def generate_pdf_mcqs(
    file: UploadFile = File(...),
    topics: str = Form("[]"),
    question_type: str = Form("both"),
    num_mcqs: int = Form(0),
    num_desc: int = Form(0),
    course_data_json: str = Form("{}"),  # Add course data as JSON string
    is_course_selection: bool = Form(False),  # Add this parameter
):
    # Update validation logic based on question_type
    if question_type not in ["mcq", "descriptive", "both"]:
        return {"status": "error", "error": "Invalid question type. Must be 'mcq', 'descriptive', or 'both'"}
    
    # Validate based on question type
    if question_type == "mcq" or question_type == "both":
        if num_mcqs <= 0:
            return {"status": "error", "error": "Number of MCQs per topic must be greater than 0 when generating MCQs"}
    
    if question_type == "descriptive" or question_type == "both":
        if num_desc < 0:
            return {"status": "error", "error": "Number of descriptive questions per topic cannot be negative"}
        if num_desc == 0 and question_type == "descriptive":
            return {"status": "error", "error": "Number of descriptive questions per topic must be greater than 0 when generating descriptive questions"}
    
    # Validate that topics were selected
    selected_topic_titles = json.loads(topics)
    if not selected_topic_titles:
        return {"status": "error", "error": "Please select at least one topic"}
    
    pdf_bytes = await file.read()
    selected_topic_titles = json.loads(topics)
    qtype = question_type.lower()
    
    # Validate that topics were selected
    if not selected_topic_titles:
        return {"status": "error", "error": "Please select at least one topic"}
    
    try:
        md5 = hashlib.md5(pdf_bytes).hexdigest()
        if STATE.get("last_pdf_hash") != md5:
            STATE["pdf_uploads"] += 1
            STATE["last_pdf_hash"] = md5

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        STATE["last_pdf_pages"] = getattr(doc, "page_count", 0)
        full_text = "\n".join([doc.load_page(p).get_text("text") or "" for p in range(doc.page_count)])
        
        # Extract all topics from PDF
        all_topics = extract_topics_from_pdf(pdf_bytes)
        
        # Filter by selected topic titles
        filtered_topics = []
        for topic in all_topics:
            if topic["title"] in selected_topic_titles:
                filtered_topics.append(topic)
        
        # If no matches found, return error
        if not filtered_topics:
            doc.close()
            return {
                "status": "error", 
                "error": f"No matching topics found. Selected topics: {selected_topic_titles}"
            }
        
        # Generate questions for each topic
        results = {}
        total_mcqs_generated = 0
        total_desc_generated = 0
        
        for topic in filtered_topics:
            title = topic["title"]
            
            # Use the topic's context or extract context from its page
            if topic.get("context"):
                context = topic["context"]
            else:
                page_num = topic["page"] - 1
                start_page = max(0, page_num - 1)
                end_page = min(doc.page_count, page_num + 2)
                context = ""
                for p in range(start_page, end_page):
                    context += doc.load_page(p).get_text("text") or ""
            
            entry = {}
            if qtype in ("mcq", "both"):
                entry["mcqs"] = generate_mcqs_ollama(
                    title, 
                    context=context[:2000], 
                    full_text=full_text, 
                    num_qs=num_mcqs
                )
                total_mcqs_generated += len(entry["mcqs"])
            else:
                entry["mcqs"] = []

            if qtype in ("descriptive", "both"):
                entry["descriptive"] = generate_descriptive_with_answers(
                    title, 
                    context=context[:2000], 
                    num_qs=num_desc
                )
                total_desc_generated += len(entry["descriptive"])
            else:
                entry["descriptive"] = []

            results[title] = entry

        doc.close()

        # Save to database
        # Instead of always saving to questions table, decide based on the flag
        pdf_filename = getattr(file, "filename", None)

        # Build output files
        df_all = build_dfs_from_questions(results)

        # CSV
        csv_bytes = df_all.to_csv(index=False).encode("utf-8")
        csv_key = hashlib.md5(csv_bytes).hexdigest()
        store_result_bytes(csv_key, csv_bytes, "questions.csv", "text/csv")

        # Excel
        excel_buf = BytesIO()
        with pd.ExcelWriter(excel_buf, engine="xlsxwriter") as writer:
            df_all.to_excel(writer, sheet_name="Questions", index=False)
        excel_buf.seek(0)
        excel_bytes = excel_buf.getvalue()
        excel_key = hashlib.md5(excel_bytes).hexdigest()
        store_result_bytes(excel_key, excel_bytes, "questions.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        # DOCX
        docx_bytes = build_docx_bytes(results)
        docx_key = hashlib.md5(docx_bytes).hexdigest()
        store_result_bytes(docx_key, docx_bytes, "questions.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Update state
        STATE["mcq_count"] = STATE.get("mcq_count", 0) + total_mcqs_generated
        STATE["desc_count"] = STATE.get("desc_count", 0) + total_desc_generated

        return {
            "status": "success",
            "results_count_topics": len(results),
            "mcqCount": total_mcqs_generated,
            "descCount": total_desc_generated,
            "download_keys": {"csv": csv_key, "excel": excel_key, "docx": docx_key},
            "pages": STATE["last_pdf_pages"],
            "global_state": {
                "pdf_uploads": STATE["pdf_uploads"],
                "last_pdf_pages": STATE["last_pdf_pages"],
                "mcq_count": STATE["mcq_count"],
                "desc_count": STATE["desc_count"]
            },
            "results": results,
            "selected_topics": [t["title"] for t in filtered_topics],
            "requested_mcqs_per_topic": num_mcqs,
            "requested_desc_per_topic": num_desc
        }

    except Exception as e:
        return {"status": "error", "error": str(e)}

@app.get("/questions")
def get_questions(search: str = None, qtype: str = None, flagged: bool = None,pdf: str = None,difficulty: str = None):
    db = SessionLocal()
    try:
        query = db.query(Question)
        
        # Filter by flagged status if provided
        if flagged is not None:
            query = query.filter(Question.flagged == flagged)
        
        if search:
            search_term = f"%{search}%"
            query = query.filter(
                Question.question.ilike(search_term) |
                Question.topic.ilike(search_term) |
                Question.option_a.ilike(search_term) |
                Question.option_b.ilike(search_term) |
                Question.option_c.ilike(search_term) |
                Question.option_d.ilike(search_term) |
                Question.answer.ilike(search_term) |
                Question.descriptive_answer.ilike(search_term)
            )
        if pdf:
            # exact match vs ilike as you prefer; using exact match here:
            query = query.filter(Question.pdf_name == pdf)
         # ADD DIFFICULTY FILTER
        if difficulty and difficulty != 'all':
            try:
                diff_int = int(difficulty)
                query = query.filter(Question.difficulty == diff_int)
            except ValueError:
                pass
        
        # Filter by question type - FIX THIS PART
        if qtype and qtype.lower() != 'all':
            query = query.filter(Question.type == qtype)
            
        questions = query.order_by(Question.created_at.desc()).all()
        
        # Convert to dict for JSON serialization
        result = []
        for q in questions:
            result.append({
                "id": q.id,
                "topic": q.topic,
                "type": q.type,
                "question": q.question,
                "option_a": q.option_a,
                "option_b": q.option_b,
                "option_c": q.option_c,
                "option_d": q.option_d,
                "answer": q.answer,
                "descriptive_answer": q.descriptive_answer,
                "difficulty": q.difficulty,
                "pdf_name":q.pdf_name,   # <-- added
                "flagged": q.flagged,
                "created_at": q.created_at.isoformat() if q.created_at else None
            })
            
        return result
        
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)
    finally:
        db.close()

# Update the flag update function to handle individual question flagging
@app.post("/update_question_flag")
async def update_question_flag(question_data: dict):
    """
    Update the flagged status of a question
    """
    db = SessionLocal()
    try:
        question_id = question_data.get('id')
        flagged = question_data.get('flagged')
        
        if not question_id:
            return {"status": "error", "error": "Question ID is required"}
        
        question = db.query(Question).filter(Question.id == question_id).first()
        if not question:
            return {"status": "error", "error": "Question not found"}
        
        # Convert to boolean to ensure consistent data type
        question.flagged = flagged
        db.commit()
        
        return {
            "status": "success", 
            "message": f"Question {question_id} flagged status updated to {flagged}",
            "question_id": question_id,
            "flagged": bool(flagged)
        }
    
    except Exception as e:
        db.rollback()
        return {"status": "error", "error": str(e)}
    finally:
        db.close()



@app.post("/save_questions_to_db")
async def save_questions_to_db_endpoint(data: dict):
    try:
        # Extract results and pdf_name from data
        results = data.get('results', {})
        pdf_name = data.get('pdf_name', None)
        
        # Generate a hash for this result set
        import json
        result_str = json.dumps(results, sort_keys=True)
        result_hash = hashlib.md5(result_str.encode()).hexdigest()
        
        # Check if already exported
        if result_hash in EXPORTED_RESULTS:
            return JSONResponse(content={
                "status": "already_exported",
                "message": "These questions have already been exported to the database",
                "timestamp": EXPORTED_RESULTS[result_hash].get('timestamp')
            })
        
        # Save to database WITH pdf_name
        save_result = save_questions_to_db(results, pdf_name=pdf_name)
        
        # Mark as exported
        if save_result.get("status") == "success":
            EXPORTED_RESULTS[result_hash] = {
                "exported": True,
                "timestamp": datetime.datetime.utcnow().isoformat(),
                "pdf_name": pdf_name
            }
        
        return JSONResponse(content=save_result)
    except Exception as e:
        return JSONResponse(content={"status": "error", "error": str(e)}, status_code=500)





import re
from random import sample



from sqlalchemy import or_, and_

@app.post("/check_export_status")
async def check_export_status(data: dict):
    """Check if results have already been exported"""
    try:
        result_hash = data.get('result_hash')
        if not result_hash:
            return {"status": "error", "error": "No result hash provided"}
        
        is_exported = result_hash in EXPORTED_RESULTS
        
        return {
            "status": "success",
            "is_exported": is_exported,
            "timestamp": EXPORTED_RESULTS.get(result_hash, {}).get('timestamp')
        }
    except Exception as e:
        return {"status": "error", "error": str(e)}
@app.get("/get_components/{course_id}")
def get_components(course_id: str):
    """Get Component values for a specific course from tblCourseEvaluationComponent"""
    try:
        import pyodbc
        
        
        
        # SQL Server Connection
        sql_conn = pyodbc.connect(
            "DRIVER={ODBC Driver 17 for SQL Server};"
            "SERVER=192.168.1.56,1433;"
            "DATABASE=ICFAISMS;"
            "UID=sa;"
            "PWD=icfai@123;"
            "TrustServerCertificate=yes;"
        )
        
        cursor = sql_conn.cursor()
        
        # Get schema for tblCourseEvaluationComponent
        cursor.execute("""
            SELECT TABLE_SCHEMA 
            FROM INFORMATION_SCHEMA.TABLES 
            WHERE TABLE_NAME = 'tblCourseEvaluationComponent'
        """)
        
        schema_result = cursor.fetchone()
        schema = schema_result[0] if schema_result else "dbo"
        full_table_name = f"{schema}.tblCourseEvaluationComponent"
        
        
        
        # Get distinct Components for this course
        query = f"""
            SELECT DISTINCT Component 
            FROM {full_table_name} 
            WHERE CourseID = ? 
            AND Component IS NOT NULL 
            ORDER BY Component
        """
        
        cursor.execute(query, course_id)
        components = [str(row[0]) for row in cursor.fetchall()]
        
        
        cursor.close()
        sql_conn.close()
        
        result = {
            "components": components,
            "course_id": course_id,
            "table_used": full_table_name
        }
        
        
        
        return result
        
    except Exception as e:
        print(f"[ERROR] Exception in get_components: {str(e)}")
        import traceback
        traceback.print_exc()
        return {"error": str(e)}

@app.get("/get_program_details/{faculty_id}/{program_id}")
def get_program_details(faculty_id: str, program_id: str):
    """Get Branch, AcademicYear, Semester, Course details with names for a specific program"""
    try:
        import pyodbc
        
        
        
        # SQL Server Connection
        sql_conn = pyodbc.connect(
            "DRIVER={ODBC Driver 17 for SQL Server};"
            "SERVER=192.168.1.56,1433;"
            "DATABASE=ICFAISMS;"
            "UID=sa;"
            "PWD=icfai@123;"
            "TrustServerCertificate=yes;"
        )
        
        cursor = sql_conn.cursor()
        
        # Get schema for tblProgBranchYearMap
        cursor.execute("""
            SELECT TABLE_SCHEMA 
            FROM INFORMATION_SCHEMA.TABLES 
            WHERE TABLE_NAME = 'tblProgBranchYearMap'
        """)
        
        schema_result = cursor.fetchone()
        schema = schema_result[0] if schema_result else "dbo"
        full_table_name = f"{schema}.tblProgBranchYearMap"
        
       
        
        # Get schemas for master tables
        def get_table_schema(table_name):
            cursor.execute(f"""
                SELECT TABLE_SCHEMA 
                FROM INFORMATION_SCHEMA.TABLES 
                WHERE TABLE_NAME = '{table_name}'
            """)
            result = cursor.fetchone()
            return result[0] if result else "dbo"
        
        branch_schema = get_table_schema('tblBranch')
        year_schema = get_table_schema('tblAcademicYear')
        semester_schema = get_table_schema('tblSemester')
        course_schema = get_table_schema('tblCourse')
        
        # First, check if the ProgramID exists
        check_query = f"""
            SELECT COUNT(*) 
            FROM {full_table_name} 
            WHERE ProgramID = ?
        """
        cursor.execute(check_query, program_id)
        count = cursor.fetchone()[0]
        
        
        # Get distinct Branches with names
        query = f"""
            SELECT DISTINCT pbym.BranchID, b.BranchName
            FROM {full_table_name} pbym
            LEFT JOIN {branch_schema}.tblBranch b ON pbym.BranchID = b.BranchID
            WHERE pbym.ProgramID = ? 
            AND pbym.BranchID IS NOT NULL 
            ORDER BY b.BranchName
        """
       
        cursor.execute(query, program_id)
        branches = [{"id": str(row[0]), "name": str(row[1]) if row[1] else str(row[0])} 
                    for row in cursor.fetchall()]
        
        
        # Get distinct Academic Years with names
        query = f"""
            SELECT DISTINCT pbym.AcademicYearID, ay.AcademicYearName
            FROM {full_table_name} pbym
            LEFT JOIN {year_schema}.tblAcademicYear ay ON pbym.AcademicYearID = ay.AcademicYearID
            WHERE pbym.ProgramID = ? 
            AND pbym.AcademicYearID IS NOT NULL
            ORDER BY ay.AcademicYearName
        """
        
        cursor.execute(query, program_id)
        academic_years = [{"id": str(row[0]), "name": str(row[1]) if row[1] else str(row[0])} 
                          for row in cursor.fetchall()]
        
        
        # Get distinct Semesters with names
        query = f"""
            SELECT DISTINCT pbym.SemesterID, s.SemesterName
            FROM {full_table_name} pbym
            LEFT JOIN {semester_schema}.tblSemester s ON pbym.SemesterID = s.SemesterID
            WHERE pbym.ProgramID = ? 
            AND pbym.SemesterID IS NOT NULL
            ORDER BY s.SemesterName
        """
        
        cursor.execute(query, program_id)
        semesters = [{"id": str(row[0]), "name": str(row[1]) if row[1] else str(row[0])} 
                     for row in cursor.fetchall()]
        
        
        # Get distinct Courses with names
        query = f"""
            SELECT DISTINCT pbym.CourseID, c.CourseName
            FROM {full_table_name} pbym
            LEFT JOIN {course_schema}.tblCourse c ON pbym.CourseID = c.CourseID
            WHERE pbym.ProgramID = ? 
            AND pbym.CourseID IS NOT NULL
            ORDER BY c.CourseName
        """
        
        cursor.execute(query, program_id)
        courses = [{"id": str(row[0]), "name": str(row[1]) if row[1] else str(row[0])} 
                   for row in cursor.fetchall()]
        
        
        cursor.close()
        sql_conn.close()
        
        result = {
            "branches": branches,
            "academic_years": academic_years,
            "semesters": semesters,
            "courses": courses,
            "program_id": program_id,
            "faculty_id": faculty_id,
            "table_used": full_table_name,
            "total_rows_for_program": count
        }
        
        
        
        return result
        
    except Exception as e:
        print(f"[ERROR] Exception in get_program_details: {str(e)}")
        import traceback
        traceback.print_exc()
        return {"error": str(e)}

@app.get("/get_faculty_filters/{faculty_id}")
def get_faculty_filters(faculty_id: str):
    """Get ProgramIDs and ProgramNames for a specific faculty from tblFacultyProgram"""
    try:
        import pyodbc

        # SQL Server Connection
        sql_conn = pyodbc.connect(
            "DRIVER={ODBC Driver 17 for SQL Server};"
            "SERVER=192.168.1.56,1433;"
            "DATABASE=ICFAISMS;"
            "UID=sa;"
            "PWD=icfai@123;"
            "TrustServerCertificate=yes;"
        )

        cursor = sql_conn.cursor()

        # Get schema for tblFacultyProgram
        cursor.execute("""
            SELECT TABLE_SCHEMA 
            FROM INFORMATION_SCHEMA.TABLES 
            WHERE TABLE_NAME = 'tblFacultyProgram'
        """)
        schema_result = cursor.fetchone()
        schema = schema_result[0] if schema_result else "dbo"
        full_table_name = f"{schema}.tblFacultyProgram"

        

        # Get schema for tblProgramMaster
        cursor.execute("""
            SELECT TABLE_SCHEMA 
            FROM INFORMATION_SCHEMA.TABLES 
            WHERE TABLE_NAME = 'tblProgramMaster'
        """)
        prog_schema_result = cursor.fetchone()
        prog_schema = prog_schema_result[0] if prog_schema_result else "dbo"
        prog_table = f"{prog_schema}.tblProgramMaster"

        # Join with tblProgramMaster to get ProgramName
        query = f"""
            SELECT DISTINCT fp.ProgramID, pm.ProgramName
            FROM {full_table_name} fp
            LEFT JOIN {prog_table} pm ON fp.ProgramID = pm.ProgramMasterID
            WHERE fp.FacultyID = ? 
              AND fp.ProgramID IS NOT NULL 
            ORDER BY pm.ProgramName
        """
        cursor.execute(query, (faculty_id,))
        programs = [{"id": str(row[0]).strip(), "name": str(row[1]).strip() if row[1] else str(row[0]).strip()} 
                    for row in cursor.fetchall()]

        cursor.close()
        sql_conn.close()

        

        return {
            "programs": programs,
            "faculty_id": faculty_id,
            "table_used": full_table_name
        }

    except Exception as e:
        print(f"[ERROR] Exception in get_faculty_filters: {str(e)}")
        import traceback
        traceback.print_exc()
        return {"error": str(e)}



@app.post("/generate_question_paper")
async def generate_question_paper(request_data: dict):
    """
    Generate a question paper with random questions based on the selected levels, types, and topics.
    """
    db = SessionLocal()
    try:
        # Extract parameters from request data
        levels = request_data.get('levels', {})
        types = request_data.get('types', {'mcq': True, 'descriptive': True})
        topics = request_data.get('topics', 'all')
        
        # Convert topics to list if it's a string
        if topics == 'all':
            selected_topics = []
        else:
            selected_topics = topics if isinstance(topics, list) else [topics]
        
        # Build query filters
        query_filters = []
        
        # Filter by question type
        type_filters = []
        if types.get('mcq', True):
            type_filters.append(Question.type == 'MCQ')
        if types.get('descriptive', True):
            type_filters.append(Question.type == 'Descriptive')
        
        if type_filters:
            query_filters.append(or_(*type_filters))
        
        # Filter by topic if specific topics are selected
        if selected_topics:
            query_filters.append(Question.topic.in_(selected_topics))
        # IMPORTANT: only approved questions
        query_filters.append(Question.flagged == True)
        
        # Apply filters to query
        query = db.query(Question)
        if query_filters:
            query = query.filter(and_(*query_filters))
        
        all_questions = query.all()
        
        # Group questions by difficulty level
        questions_by_level = {1: [], 2: [], 3: [], 4: [], 5: []}
        
        for q in all_questions:
            if q.difficulty and q.difficulty.isdigit():
                level = int(q.difficulty)
                if 1 <= level <= 5:
                    questions_by_level[level].append(q)
        
        # Create a paper by selecting random questions from each level
        question_paper = []
        total_selected = 0
        level_summary = {}
        
        for level, count in levels.items():
            level = int(level)  # Ensure level is integer
            if count > 0 and level in questions_by_level:
                available_questions = questions_by_level[level]
                if available_questions:
                    num_to_select = min(count, len(available_questions))
                    selected_questions = sample(available_questions, num_to_select)
                    question_paper.extend(selected_questions)
                    total_selected += num_to_select
                    level_summary[level] = num_to_select
                else:
                    level_summary[level] = 0
        
        # Return the selected question paper data
        paper_data = []
        for q in question_paper:
            # Clean the options to remove answer and difficulty info
            def clean_option(option_text):
                if not option_text:
                    return option_text
                
                # Remove "Answer: X Difficulty: Y" patterns from options
                option_text = re.sub(r'\s*Answer:\s*[A-D]\s*Difficulty:\s*\d\s*$', '', option_text, flags=re.IGNORECASE)
                option_text = re.sub(r'\s*Difficulty:\s*\d\s*Answer:\s*[A-D]\s*$', '', option_text, flags=re.IGNORECASE)
                
                # Remove standalone patterns
                option_text = re.sub(r'\s*Answer:\s*[A-D]\s*$', '', option_text, flags=re.IGNORECASE)
                option_text = re.sub(r'\s*Difficulty:\s*\d\s*$', '', option_text, flags=re.IGNORECASE)
                
                # Final cleanup
                option_text = re.sub(r'[\.\s]*$', '', option_text).strip()
                return option_text

            # Add sanitized question to the result
            question_dict = {
                "id": q.id,
                "topic": q.topic,
                "type": q.type,
                "question": q.question.strip(),
                "option_a": clean_option(q.option_a),
                "option_b": clean_option(q.option_b),
                "option_c": clean_option(q.option_c),
                "option_d": clean_option(q.option_d),
                "flagged": q.flagged,
                "difficulty": q.difficulty
            }

            paper_data.append(question_dict)
        
        return {
            "status": "success", 
            "questions": paper_data,
            "total_selected": total_selected,
            "level_summary": level_summary,
            "filters_applied": {
                "levels": levels,
                "types": types,
                "topics": selected_topics if selected_topics else "all"
            },
            "message": f"Generated paper with {total_selected} questions"
        }
    
    except Exception as e:
        return {"status": "error", "error": str(e)}
    finally:
        db.close()

@app.post("/update_question")
async def update_question(question_data: dict):
    """
    Update any field of a question
    """
    db = SessionLocal()
    try:
        question_id = question_data.get('id')
        updates = question_data.get('updates', {})
        
        if not question_id:
            return {"status": "error", "error": "Question ID is required"}
        
        question = db.query(Question).filter(Question.id == question_id).first()
        if not question:
            return {"status": "error", "error": "Question not found"}
        
        # Update allowed fields
        allowed_fields = ['topic', 'question', 'option_a', 'option_b', 'option_c', 'option_d', 
                         'answer', 'descriptive_answer', 'difficulty', 'flagged']
        
        for field, value in updates.items():
            if field in allowed_fields and hasattr(question, field):
                setattr(question, field, value)
        
        db.commit()
        
        return {
            "status": "success", 
            "message": f"Question {question_id} updated successfully",
            "question_id": question_id,
            "updates": updates
        }
    
    except Exception as e:
        db.rollback()
        return {"status": "error", "error": str(e)}
    finally:
        db.close()




@app.post("/bulk_update_flags")
async def bulk_update_flags(bulk_data: dict):
    """
    Update flagged status for multiple questions at once
    """
    db = SessionLocal()
    try:
        question_updates = bulk_data.get('question_updates', [])
        
        if not question_updates:
            return {"status": "error", "error": "No question updates provided"}
        
        updated_count = 0
        for update in question_updates:
            question_id = update.get('id')
            flagged = update.get('flagged')
            
            if question_id is not None:
                question = db.query(Question).filter(Question.id == question_id).first()
                if question:
                    question.flagged = flagged
                    updated_count += 1
        
        db.commit()
        
        return {
            "status": "success", 
            "message": f"Updated flagged status for {updated_count} questions",
            "updated_count": updated_count
        }
    
    except Exception as e:
        db.rollback()
        return {"status": "error", "error": str(e)}
    finally:
        db.close()

@app.post("/faculty_login")
async def faculty_login(login_data: dict):
    """Authenticate faculty using FacultyID from tblFacultyProgram"""
    try:
        import pyodbc
        
        faculty_id = login_data.get('faculty_id')
        
        if not faculty_id:
            return {"status": "error", "error": "Faculty User ID is required"}
        
        # SQL Server Connection
        sql_conn = pyodbc.connect(
            "DRIVER={ODBC Driver 17 for SQL Server};"
            "SERVER=192.168.1.56,1433;"
            "DATABASE=ICFAISMS;"
            "UID=sa;"
            "PWD=icfai@123;"
            "TrustServerCertificate=yes;"
        )
        
        cursor = sql_conn.cursor()
        
        # Check the schema for tblFacultyProgram table
        cursor.execute("""
            SELECT TABLE_SCHEMA 
            FROM INFORMATION_SCHEMA.TABLES 
            WHERE TABLE_NAME = 'tblFacultyProgram'
        """)
        
        schema_result = cursor.fetchone()
        schema = schema_result[0] if schema_result else "dbo"
        full_table_name = f"{schema}.tblFacultyProgram"
        
        
        
        # Check if faculty exists in tblFacultyProgram using FacultyID
        query = f"""
            SELECT DISTINCT FacultyID 
            FROM {full_table_name} 
            WHERE FacultyID = ?
        """
        
        
        cursor.execute(query, faculty_id)
        faculty = cursor.fetchone()
        
        if not faculty:
            cursor.close()
            sql_conn.close()
            return {"status": "error", "error": f"Invalid Faculty User ID: {faculty_id}"}
        
        # Try to get faculty name from tblEmployee or tblUsers
        faculty_name = f"Faculty {faculty_id}"
        email = ""
        
        try:
            # Check tblEmployee table
            cursor.execute("""
                SELECT TABLE_SCHEMA 
                FROM INFORMATION_SCHEMA.TABLES 
                WHERE TABLE_NAME = 'tblEmployee'
            """)
            emp_schema_result = cursor.fetchone()
            emp_schema = emp_schema_result[0] if emp_schema_result else "dbo"
            
            # Try multiple possible column names for faculty identification
            cursor.execute(f"""
                SELECT EmployeeName, Email 
                FROM {emp_schema}.tblEmployee 
                WHERE UserID=?
            """, faculty_id)
            
            emp_result = cursor.fetchone()
            if emp_result:
                faculty_name = emp_result[0] if emp_result[0] else faculty_name
                email = emp_result[1] if emp_result[1] else ""
            else:
                # Try tblUsers table
                cursor.execute("""
                    SELECT TABLE_SCHEMA 
                    FROM INFORMATION_SCHEMA.TABLES 
                    WHERE TABLE_NAME = 'tblUsers'
                """)
                user_schema_result = cursor.fetchone()
                user_schema = user_schema_result[0] if user_schema_result else "dbo"
                
                cursor.execute(f"""
                    SELECT UserName, Email 
                    FROM {user_schema}.tblUsers 
                    WHERE UserID = ? 
                       OR FacultyUserID = ? 
                       OR FacultyID = ?
                """, faculty_id, faculty_id)
                
                user_result = cursor.fetchone()
                if user_result:
                    faculty_name = user_result[0] if user_result[0] else faculty_name
                    email = user_result[1] if user_result[1] else ""
        except Exception as e:
            print(f"Note: Could not fetch faculty details: {e}")
        
        cursor.close()
        sql_conn.close()
        
        return {
            "status": "success",
            "faculty_id": faculty_id,
            "faculty_name": faculty_name,
            "email": email,
            "authenticated": True,
            "table_used": full_table_name
        }
        
    except Exception as e:
        print(f"Login error: {e}")
        return {"status": "error", "error": str(e)}

@app.get("/get_program_outcomes/{program_id}")
def get_program_outcomes(program_id: str):
    """Get Program Outcomes for a specific program"""
    try:
        import pyodbc
        
        
        
        # SQL Server Connection
        sql_conn = pyodbc.connect(
            "DRIVER={ODBC Driver 17 for SQL Server};"
            "SERVER=192.168.1.56,1433;"
            "DATABASE=ICFAISMS;"
            "UID=sa;"
            "PWD=icfai@123;"
            "TrustServerCertificate=yes;"
        )
        
        cursor = sql_conn.cursor()
        
        # Get schema for tblProgramOutcome
        cursor.execute("""
            SELECT TABLE_SCHEMA 
            FROM INFORMATION_SCHEMA.TABLES 
            WHERE TABLE_NAME = 'tblProgramOutcome'
        """)
        
        schema_result = cursor.fetchone()
        schema = schema_result[0] if schema_result else "dbo"
        full_table_name = f"{schema}.tblProgramOutcome"
        
        
        
        # First, let's see what columns are available
        cursor.execute(f"""
            SELECT TOP 1 * FROM {full_table_name}
        """)
        columns = [column[0] for column in cursor.description]
        
        
        # Get Program Outcomes - using ProgramOutcomeID instead of POID
        query = f"""
            SELECT DISTINCT ProgramOutcomeID, PODescription
            FROM {full_table_name}
            WHERE ProgramID = ?
            AND PODescription IS NOT NULL
            ORDER BY ProgramOutcomeID
        """
        
       
        cursor.execute(query, program_id)
        outcomes = [{"id": str(row[0]), "description": str(row[1])} 
                   for row in cursor.fetchall()]
        
        cursor.close()
        sql_conn.close()
        
        
        
        return {
            "program_outcomes": outcomes,
            "program_id": program_id,
            "table_used": full_table_name,
            "count": len(outcomes)
        }
        
    except Exception as e:
        print(f"[ERROR] Exception in get_program_outcomes: {str(e)}")
        import traceback
        traceback.print_exc()
        return {"error": str(e), "program_outcomes": []}


@app.get("/get_course_outcomes/{course_id}")
def get_course_outcomes(course_id: str):
    """Get Course Outcomes for a specific course"""
    try:
        import pyodbc
        
       
        
        # SQL Server Connection
        sql_conn = pyodbc.connect(
            "DRIVER={ODBC Driver 17 for SQL Server};"
            "SERVER=192.168.1.56,1433;"
            "DATABASE=ICFAISMS;"
            "UID=sa;"
            "PWD=icfai@123;"
            "TrustServerCertificate=yes;"
        )
        
        cursor = sql_conn.cursor()
        
        # Get schema for tblCourseOutcome
        cursor.execute("""
            SELECT TABLE_SCHEMA 
            FROM INFORMATION_SCHEMA.TABLES 
            WHERE TABLE_NAME = 'tblCourseOutcome'
        """)
        
        schema_result = cursor.fetchone()
        schema = schema_result[0] if schema_result else "dbo"
        full_table_name = f"{schema}.tblCourseOutcome"
        
        
        
        # First, let's see what columns are available
        cursor.execute(f"""
            SELECT TOP 1 * FROM {full_table_name}
        """)
        columns = [column[0] for column in cursor.description]
        
        
        # Get Course Outcomes - using CourseOutcomeID instead of COID
        query = f"""
            SELECT DISTINCT CourseOutcomeID, CODescription
            FROM {full_table_name}
            WHERE CourseID = ?
            AND CODescription IS NOT NULL
            ORDER BY CourseOutcomeID
        """
        
        
        cursor.execute(query, course_id)
        outcomes = [{"id": str(row[0]), "description": str(row[1])} 
                   for row in cursor.fetchall()]
        
        cursor.close()
        sql_conn.close()
        
        
        
        return {
            "course_outcomes": outcomes,
            "course_id": course_id,
            "table_used": full_table_name,
            "count": len(outcomes)
        }
        
    except Exception as e:
        print(f"[ERROR] Exception in get_course_outcomes: {str(e)}")
        import traceback
        traceback.print_exc()
        return {"error": str(e), "course_outcomes": []}
    
### extract topics from PDF
def extract_topics_from_pdf(pdf_bytes: bytes) -> List[Dict]:
    """
    Extract meaningful topics/headings from PDF using multiple strategies.
    Returns list of topics with their content.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    topics = []
    
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        
        # Strategy 1: Extract text blocks with larger fonts (likely headings)
        blocks = page.get_text("dict")["blocks"]
        
        for block in blocks:
            if "lines" in block:
                for line in block["lines"]:
                    if line["spans"]:
                        # Get the largest font size in this line
                        max_font_size = max(span["size"] for span in line["spans"])
                        text = "".join(span["text"] for span in line["spans"]).strip()
                        
                        # Criteria for identifying headings/topics
                        if (max_font_size >= 11 and  # Larger than body text
                            len(text) > 3 and len(text) < 200 and  # Reasonable length
                            text and not text.isspace() and
                            not text.startswith(('http://', 'https://', 'www.')) and
                            text[0].isalpha() and  # Starts with letter
                            not text.endswith(('.', ',', ';', ':'))):  # Not sentence endings
                            
                            # Clean the text
                            clean_text = re.sub(r'\s+', ' ', text).strip()
                            
                            # Skip common non-topic text
                            skip_patterns = [
                                "page", "copyright", "©", "confidential",
                                "draft", "table of contents", "contents",
                                "index", "appendix", "references", "bibliography",
                                "acknowledgment", "abstract", "chapter", "part",
                                "figure", "table", "list of"
                            ]
                            
                            if not any(pattern in clean_text.lower() for pattern in skip_patterns):
                                # Get surrounding context for better topic understanding
                                context_start = max(0, page_num - 1)
                                context_end = min(doc.page_count, page_num + 2)
                                context = ""
                                for p in range(context_start, context_end):
                                    context += doc.load_page(p).get_text("text") or ""
                                
                                topics.append({
                                    "title": clean_text,
                                    "page": page_num + 1,
                                    "font_size": max_font_size,
                                    "context": context[:1000],  # First 1000 chars for context
                                    "confidence": min(1.0, max_font_size / 16.0)
                                })
        
        # Strategy 2: Look for numbered sections (e.g., "1. Introduction", "2.1 Background")
        text = page.get_text("text")
        numbered_patterns = [
            r'^\s*(\d+\.\d*)\s+(.+?)$',  # 1.1 Topic
            r'^\s*([IVXLCDM]+\.)\s+(.+?)$',  # I. Topic, II. Topic
            r'^\s*([A-Z]\.)\s+(.+?)$',  # A. Topic, B. Topic
            r'^\s*(\d+\))\s+(.+?)$',  # 1) Topic, 2) Topic
        ]
        
        for line in text.split('\n'):
            line = line.strip()
            for pattern in numbered_patterns:
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    topic_text = match.group(2) if len(match.groups()) > 1 else match.group(1)
                    if len(topic_text) > 5 and len(topic_text) < 150:
                        topics.append({
                            "title": topic_text.strip(),
                            "page": page_num + 1,
                            "font_size": 12.0,
                            "context": "",
                            "confidence": 0.8,
                            "type": "numbered"
                        })
                    break
    
    doc.close()
    
    # Remove duplicates and similar topics
    unique_topics = []
    seen_titles = set()
    
    for topic in topics:
        title_lower = topic["title"].lower()
        
        # Check if similar topic already exists
        is_duplicate = False
        for seen in seen_titles:
            # Fuzzy matching for similar topics
            if (title_lower == seen or
                title_lower in seen or
                seen in title_lower or
                fuzz.ratio(title_lower, seen) > 80):  # Using fuzzy matching
                is_duplicate = True
                break
        
        if not is_duplicate and len(title_lower) > 3:
            seen_titles.add(title_lower)
            unique_topics.append(topic)
    
    # Sort by page number, then by font size
    unique_topics.sort(key=lambda x: (x["page"], -x.get("font_size", 0)))
    
    return unique_topics[:30]  # Return top 30 topics
def estimate_generation_time(num_topics: int, num_mcqs_per_topic: int, num_desc_per_topic: int, question_type: str) -> Dict:
    """
    Estimate time needed for question generation.
    Based on average processing times:
    - Topic analysis: 2 seconds per topic
    - MCQ generation: 3 seconds per question (Ollama)
    - Descriptive generation: 4 seconds per question (Ollama)
    """
    
    # Base processing time (PDF loading, topic extraction, etc.)
    base_time_seconds = 10
    
    # Time per topic for context extraction
    topic_analysis_time = num_topics * 2
    
    # Calculate question generation time
    mcq_time = 0
    desc_time = 0
    
    if question_type in ["mcq", "both"]:
        total_mcqs = num_topics * num_mcqs_per_topic
        mcq_time = total_mcqs * 10
    
    if question_type in ["descriptive", "both"]:
        total_desc = num_topics * num_desc_per_topic
        desc_time = total_desc * 10
    
    total_seconds = base_time_seconds + topic_analysis_time + mcq_time + desc_time
    
    # Convert to human-readable format
    if total_seconds < 60:
        time_str = f"About {total_seconds} seconds"
    elif total_seconds < 3600:
        minutes = total_seconds // 60
        seconds = total_seconds % 60
        time_str = f"About {minutes} minute{'s' if minutes > 1 else ''}"
        if seconds > 0 and minutes < 5:
            time_str += f" {seconds} seconds"
    else:
        hours = total_seconds // 3600
        minutes = (total_seconds % 3600) // 60
        time_str = f"About {hours} hour{'s' if hours > 1 else ''}"
        if minutes > 0:
            time_str += f" {minutes} minute{'s' if minutes > 1 else ''}"
    
    return {
        "estimated_seconds": total_seconds,
        "estimated_time": time_str,
        "details": {
            "num_topics": num_topics,
            "total_mcqs": num_topics * num_mcqs_per_topic if question_type in ["mcq", "both"] else 0,
            "total_descriptive": num_topics * num_desc_per_topic if question_type in ["descriptive", "both"] else 0,
            "question_type": question_type
        }
    }

@app.post("/estimate_generation_time")
async def get_generation_time_estimate(data: dict):
    """
    Estimate time for question generation.
    Expected data: {
        "num_topics": int,
        "num_mcqs": int,
        "num_desc": int,
        "question_type": str
    }
    """
    try:
        num_topics = data.get("num_topics", 0)
        num_mcqs = data.get("num_mcqs", 0)
        num_desc = data.get("num_desc", 0)
        question_type = data.get("question_type", "both")
        
        estimate = estimate_generation_time(num_topics, num_mcqs, num_desc, question_type)
        
        return {
            "status": "success",
            "estimate": estimate
        }
    except Exception as e:
        return {"status": "error", "error": str(e)}
@app.get("/get_combined_filters")
async def get_combined_filters():
    db = SessionLocal()
    try:
        # distinct pdf names (skip null/empty)
        rows = db.query(Question.pdf_name).distinct().all()
        pdf_names = [r[0] for r in rows if r and r[0]]
        # keep existing payload structure you use on frontend; here minimal
        return {
            "pdf_names": pdf_names,
            # you can also return course_ids, academic_years, semesters if you collect them
            "course_ids": [], "academic_years": [], "semesters": []
        }
    except Exception as e:
        return {"error": str(e)}
    finally:
        db.close()

@app.post("/save_questions_to_course_db")
async def save_questions_to_course_db_endpoint(request_data: dict):
    try:
        print("📥 Received request to save to course_questions")
        
        # Extract course data and results
        results = request_data.get("results", {})
        pdf_name = request_data.get("pdf_name", "")
        course_data = request_data.get("course_data", {})
        
        print(f"📊 Results keys: {list(results.keys())}")
        print(f"📄 PDF name: {pdf_name}")
        print(f"🎓 Course data: {course_data}")
        
        # Generate a hash for this result set
        import json
        result_str = json.dumps({
            "results": results,
            "course_data": course_data
        }, sort_keys=True)
        result_hash = hashlib.md5(result_str.encode()).hexdigest()
        
        # Check if already exported
        if result_hash in EXPORTED_RESULTS:
            return JSONResponse(content={
                "status": "already_exported",
                "message": "These questions have already been exported to the course database",
                "timestamp": EXPORTED_RESULTS[result_hash].get('timestamp')
            })
        
        # Save to course questions table
        result = save_questions_to_course_db(results, pdf_name, course_data)
        
        # Mark as exported if successful
        if result.get("status") == "success":
            EXPORTED_RESULTS[result_hash] = {
                "exported": True,
                "timestamp": datetime.datetime.utcnow().isoformat(),
                "pdf_name": pdf_name,
                "course_data": course_data
            }
        
        return JSONResponse(content=result)
    except Exception as e:
        print("❌ Error in save_questions_to_course_db_endpoint:", e)
        import traceback
        traceback.print_exc()
        return JSONResponse(content={"status": "error", "error": str(e)}, status_code=500)


@app.get("/course_questions")
def get_course_questions(
    search: str = None, 
    qtype: str = None, 
    flagged: bool = None,
    pdf: str = None,
    difficulty: str = None,
    program_id: str = None,
    course_id: str = None
):
    db = SessionLocal()
    try:
        query = db.query(CourseQuestion)
        
        
        # Filter by flagged status if provided
        if flagged is not None:
            query = query.filter(CourseQuestion.flagged == flagged)
        
        if search:
            search_term = f"%{search}%"
            query = query.filter(
                CourseQuestion.question.ilike(search_term) |
                CourseQuestion.topic.ilike(search_term) |
                CourseQuestion.option_a.ilike(search_term) |
                CourseQuestion.option_b.ilike(search_term) |
                CourseQuestion.option_c.ilike(search_term) |
                CourseQuestion.option_d.ilike(search_term) |
                CourseQuestion.answer.ilike(search_term) |
                CourseQuestion.descriptive_answer.ilike(search_term)
            )
        
        if pdf and pdf != 'all':
            query = query.filter(CourseQuestion.pdf_name == pdf)
        
        if difficulty and difficulty != 'all':
            try:
                diff_int = int(difficulty)
                query = query.filter(CourseQuestion.difficulty == diff_int)
            except ValueError:
                pass
        
        # Filter by question type
        if qtype and qtype.lower() != 'all':
            query = query.filter(CourseQuestion.type == qtype)
            
        # Filter by program and course if provided
        if program_id:
            query = query.filter(CourseQuestion.program_id == program_id)
        
        if course_id:
            query = query.filter(CourseQuestion.course_id == course_id)
        
        questions = query.order_by(CourseQuestion.created_at.desc()).all()
        
        
        # Convert to dict for JSON serialization
        result = []
        for q in questions:
            result.append({
                "id": q.id,
                "topic": q.topic,
                "type": q.type,
                "question": q.question,
                "option_a": q.option_a,
                "option_b": q.option_b,
                "option_c": q.option_c,
                "option_d": q.option_d,
                "answer": q.answer,
                "descriptive_answer": q.descriptive_answer,
                "difficulty": q.difficulty,
                "po_id": q.po_id,
                "co_id": q.co_id,
                "pdf_name": q.pdf_name,
                "program_id": q.program_id,
                "branch_id": q.branch_id,
                "academic_year_id": q.academic_year_id,
                "semester_id": q.semester_id,
                "course_id": q.course_id,
                "component": q.component,
                "flagged": q.flagged,
                "created_at": q.created_at.isoformat() if q.created_at else None
            })
        
        return result
        
    except Exception as e:
        print(f"❌ Error in get_course_questions: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(content={"error": str(e)}, status_code=500)
    finally:
        db.close()
@app.post("/update_course_question")
async def update_course_question(question_data: dict):
    """
    Update any field of a course question
    """
    db = SessionLocal()
    try:
        question_id = question_data.get('id')
        updates = question_data.get('updates', {})
        
        if not question_id:
            return {"status": "error", "error": "Question ID is required"}
        
        question = db.query(CourseQuestion).filter(CourseQuestion.id == question_id).first()
        if not question:
            return {"status": "error", "error": "Question not found"}
        
        # Update allowed fields
        allowed_fields = ['topic', 'question', 'option_a', 'option_b', 'option_c', 'option_d', 
                         'answer', 'descriptive_answer', 'difficulty', 'flagged', 'po_id', 'co_id']
        
        for field, value in updates.items():
            if field in allowed_fields and hasattr(question, field):
                setattr(question, field, value)
        
        db.commit()
        
        return {
            "status": "success", 
            "message": f"Course question {question_id} updated successfully",
            "question_id": question_id,
            "updates": updates
        }
    
    except Exception as e:
        db.rollback()
        return {"status": "error", "error": str(e)}
    finally:
        db.close()
@app.get("/get_available_question_counts")
def get_available_question_counts():
    """Get counts of available approved questions by difficulty level"""
    db = SessionLocal()
    try:
        # Get all approved questions
        approved_questions = db.query(Question).filter(Question.flagged == True).all()
        
        # Count by difficulty
        counts = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
        
        for question in approved_questions:
            
            if question.difficulty:
                try:
                    # Convert difficulty to integer
                    diff_str = str(question.difficulty).strip()
                    
                    # Extract number if it's in string format
                    import re
                    numbers = re.findall(r'\d+', diff_str)
                    if numbers:
                        level = int(numbers[0])
                    else:
                        # Try direct conversion
                        level = int(float(diff_str))
                    
                    # # Ensure level is between 1 and 5
                    if 1 <= level <= 5:
                        counts[level] = counts.get(level, 0) + 1
                        
                    else:
                        print(f"WARNING: Difficulty {level} out of range for question {question.id}")
                        
                except (ValueError, TypeError) as e:
                    print(f"WARNING: Could not parse difficulty '{question.difficulty}' for question {question.id}: {e}")
                    continue
        
        
        return {
            "status": "success",
            "counts": counts,
            "total": len(approved_questions),
        }
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"status": "error", "error": str(e)}
    finally:
        db.close()


# --- VIDEO / TRANSCRIPTION / SUMMARY / MCQ FLOW (paste into backend.py) ---
# required imports at top of file (if not already present)
# --- VIDEO / TRANSCRIPTION / SUMMARY / MCQ FLOW (paste into backend.py) ---
# required imports at top of file (if not already present)
import nltk
from nltk.tokenize import sent_tokenize
try:
    nltk.download('punkt', quiet=True)
except Exception:
    pass

# optional libs flags
try:
    import whisper
    _HAS_WHISPER = True
except Exception:
    _HAS_WHISPER = False

try:
    from moviepy.editor import VideoFileClip
    _HAS_MOVIEPY = True
except Exception:
    _HAS_MOVIEPY = False

# summarizer config (BART chunking)
CHUNK_WORDS = 800
SUMMARIZER_MODEL = "facebook/bart-large-cnn"
SUMMARY_MIN_LENGTH = 30

# Local summarizer via transformers (optional, heavy)
def split_transcript_into_chunks_by_words(transcript: str, chunk_words: int = CHUNK_WORDS):
    sentences = sent_tokenize(transcript)
    chunks, current, current_words = [], [], 0
    for s in sentences:
        wcount = len(s.split())
        if current_words + wcount > chunk_words and current:
            chunks.append(" ".join(current))
            current, current_words = [s], wcount
        else:
            current.append(s)
            current_words += wcount
    if current:
        chunks.append(" ".join(current))
    return chunks

def summarizer_pipeline(model_name=SUMMARIZER_MODEL):
    try:
        from transformers import pipeline
        return pipeline("summarization", model=model_name, device=-1)  # CPU
    except Exception:
        return None

def summarize_chunks(chunks, summarizer):
    summaries = []
    for c in chunks:
        if summarizer:
            try:
                out = summarizer(c, max_length=400, min_length=100, do_sample=False)
                summary_text = out[0]['summary_text'].strip()
            except Exception:
                summary_text = " ".join(c.split()[:SUMMARY_MIN_LENGTH])
        else:
            # fallback: truncate
            summary_text = " ".join(c.split()[:SUMMARY_MIN_LENGTH])
        summaries.append(summary_text)
    return summaries

def combine_and_summarize_summaries(summaries):
    if not summaries:
        return ""
    return "\n\n".join(summaries)

def summarize_transcript_with_bart(transcript: str):
    """
    Try to summarize transcript using local BART in chunks; if local summarizer not available,
    return empty chunks and caller should fallback to Ollama summarizer with summarize_text().
    """
    if not transcript or not transcript.strip():
        return {"overall": "", "chunks": []}
    chunks = split_transcript_into_chunks_by_words(transcript, CHUNK_WORDS)
    summarizer = summarizer_pipeline(SUMMARIZER_MODEL)
    if summarizer is None:
        # signal to caller that local summarizer isn't available
        return {"overall": "", "chunks": []}
    chunk_summaries = summarize_chunks(chunks, summarizer)
    overall_summary = combine_and_summarize_summaries(chunk_summaries)
    return {"overall": overall_summary, "chunks": chunk_summaries}

# Robust MCQ parser (accepts many model output formats)
def parse_mcqs_freeform(output: str) -> List[Dict]:
    mcqs = []
    if not output:
        return mcqs
    raw_lines = [ln.rstrip() for ln in output.splitlines() if ln.strip()]
    # drop very generic intro / header-only lines
    lines = []
    for ln in raw_lines:
        if re.search(r"(here are|multiple[-\s]?choice questions|based on the summary|based on the topic|following questions|the following)", ln, re.I):
            continue
        if re.match(r'^\s*(?:question|q)\s*\d+\b[:.\s-]*$', ln, re.I):
            continue
        lines.append(ln.strip())

    i = 0
    while i < len(lines):
        ln = lines[i]
        # skip stray option lines until we find a question
        if re.match(r'^[A-D][\)\.\-:]\s+', ln, re.I):
            i += 1
            continue
        question_text = re.sub(r'^\s*(?:q|question)\s*\d+\s*[:.\-\)]*\s*', '', ln, flags=re.I).strip()
        if len(question_text) < 3:
            i += 1
            continue
        # collect options
        opts = []
        opt_map = {}
        j = i + 1
        while j < len(lines) and len(opts) < 4:
            if re.match(r'^[A-D][\)\.\-:]\s+', lines[j], re.I):
                m = re.match(r'^([A-D])[\)\.\-:]\s*(.*)$', lines[j], re.I)
                if m:
                    label = m.group(1).upper()
                    text = m.group(2).strip()
                    formatted = f"{label}. {text}"
                    opts.append(formatted)
                    opt_map[label] = formatted
                else:
                    opts.append(lines[j].strip())
                j += 1
            else:
                break
        # look ahead for Answer:
        answer = ""
        look_end = min(len(lines), j + 6)
        for k in range(j, look_end):
            candidate = lines[k].strip()
            m_ans = re.match(r'(?i)^\s*(?:answer|correct)[:\s\-]*\(?\s*([A-D])\s*\)?', candidate)
            if m_ans:
                answer = m_ans.group(1).upper()
                break
            m_single = re.match(r'^\s*([A-D])[\)\.\s]*$', candidate, re.I)
            if m_single:
                answer = m_single.group(1).upper()
                break
        if answer and answer not in opt_map:
            answer = ""  # validate
        if question_text and len(opts) >= 2:
            mcqs.append({"question": question_text, "options": opts, "answer": answer})
        i = j if j > i else i + 1
    return mcqs
# whisper-based transcription (uses whisper library, raises if not installed)
def split_audio(audio_path: str, chunk_length_sec: int = 300):
    try:
        from pydub import AudioSegment
    except Exception:
        return [audio_path]
    import wave, contextlib
    with contextlib.closing(wave.open(audio_path, 'rb')) as wf:
        rate = wf.getframerate()
        n_frames = wf.getnframes()
        total_sec = n_frames / float(rate)
    if total_sec <= chunk_length_sec:
        return [audio_path]
    audio = AudioSegment.from_wav(audio_path)
    chunk_files = []
    for start_ms in range(0, len(audio), chunk_length_sec * 1000):
        chunk = audio[start_ms:start_ms + chunk_length_sec * 1000]
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
        chunk.export(tmp.name, format="wav")
        chunk_files.append(tmp.name)
    return chunk_files

def transcribe_video_bytes(video_bytes: bytes, whisper_model_name: str = "small") -> str:
    if not _HAS_WHISPER or not _HAS_MOVIEPY:
        raise RuntimeError("Whisper or moviepy not available on server.")
    # write video to temp file
    vf = tempfile.NamedTemporaryFile(delete=False, suffix=".mp4")
    vf.write(video_bytes); vf.flush(); vf.close()
    audio_path = None
    try:
        clip = VideoFileClip(vf.name)
        af = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
        audio_path = af.name
        clip.audio.write_audiofile(audio_path, logger=None)
        clip.close()
        chunk_files = split_audio(audio_path)
        model = whisper.load_model(whisper_model_name)
        full_text = ""
        for c in chunk_files:
            res = model.transcribe(c)
            text = res.get("text", "").strip()
            if text:
                full_text += text + " "
            try:
                if c != audio_path and os.path.exists(c):
                    os.remove(c)
            except Exception:
                pass
        return full_text.strip()
    finally:
        try:
            if os.path.exists(vf.name): os.remove(vf.name)
        except Exception:
            pass
        try:
            if audio_path and os.path.exists(audio_path): os.remove(audio_path)
        except Exception:
            pass

# generate MCQs from summary (reuse existing function if present)
def generate_mcqs_from_summary_local(summary: str, num_qs: int = 0, model: str = MODEL):
    # Reuse the same approach as your Streamlit function generate_mcqs_from_summary
    prompt = f"""
Generate exactly {num_qs} distinct multiple-choice questions that cover the following summary.
For each question include:
- Exactly 4 labeled options A) B) C) D) 
- A single-letter answer line like: Answer: <A/B/C/D>

Use exactly this format; do not add extra commentary or code fences.

Q1. <question text>
A) <option A>
B) <option B>
C) <option C>
D) <option D>
Answer: <A/B/C/D>

Summary:
{summary}
"""
    out = call_ollama(prompt, model=model, timeout=600)
    if out.startswith("OLLAMA_ERROR"):
        return [{"question": out, "options": [], "answer": ""}]
    return parse_mcqs_freeform(out)

# Endpoint: transcribe -> summarize (video)
@app.post("/transcribe_video")
async def transcribe_video(file: UploadFile = File(...), whisper_model: str = Form("small")):
    """
    Accepts a video file and returns transcript + summary.
    If local BART summarizer (transformers) is available it will be used; otherwise Ollama summarization used.
    """
    video_bytes = await file.read()
    try:
        # Transcribe (Whisper)
        if not _HAS_WHISPER or not _HAS_MOVIEPY:
            return {"status": "error", "error": "Transcription requires whisper and moviepy installed on server."}
        # update unique-video counter
        try:
            md5 = hashlib.md5(video_bytes).hexdigest()
            if STATE.get("last_video_hash") != md5:
                STATE["video_uploads"] = STATE.get("video_uploads", 0) + 1
                STATE["last_video_hash"] = md5
        except Exception:
            pass
        transcript = transcribe_video_bytes(video_bytes, whisper_model_name=whisper_model)
        # Try local BART summarizer first
        summ = summarize_transcript_with_bart(transcript)
        if not summ["overall"]:
            # fallback: use Ollama summarizer (summarize_text uses Ollama)
            overall = summarize_text(transcript, model=MODEL, max_words=200)
            return {"status": "success", "transcript": transcript, "summary": overall, "chunks": summ["chunks"]}
        return {"status": "success", "transcript": transcript, "summary": summ["overall"], "chunks": summ["chunks"],"global_state": {
                "video_uploads": STATE.get("video_uploads", 0),}}
    except Exception as e:
        return {"status": "error", "error": str(e)}

# Endpoint: generate MCQs (from summary or from video file)
@app.post("/generate_video_mcqs")
async def generate_video_mcqs(
    file: UploadFile = File(None),
    summary: str = Form(""),
    question_type: str = Form("both"),   # "mcq", "descriptive", "both"
    num_qs: int = Form(0),
    whisper_model: str = Form("small")
):
    """
    Generate MCQs (and optionally descriptive questions) from a provided summary string,
    or from an uploaded video file (which will be transcribed & summarized).
    Returns per-request counts and download keys.
    """
    qtype = (question_type or "both").lower()
    summary_text = summary or ""
    try:
        # If file provided and summary empty, transcribe & summarize first
        if file is not None and not summary_text:
            if not _HAS_WHISPER or not _HAS_MOVIEPY:
                return {"status": "error", "error": "Transcription requires whisper and moviepy installed on server."}
            video_bytes = await file.read()
            transcript = transcribe_video_bytes(video_bytes, whisper_model_name=whisper_model)
            # try local BART
            summ = summarize_transcript_with_bart(transcript)
            if summ["overall"]:
                summary_text = summ["overall"]
                chunk_summaries = summ["chunks"]
            else:
                # fallback to Ollama
                summary_text = summarize_text(transcript, model=MODEL, max_words=200)
                chunk_summaries = summ["chunks"]
        elif summary_text:
            chunk_summaries = []
        else:
            return {"status": "error", "error": "No summary or file provided."}

        produce_mcq = (qtype in ("mcq", "both"))
        produce_desc = (qtype in ("descriptive", "both"))

        results = {}
        # We'll treat this as single topic "Video Summary"
        if produce_mcq:
            mcqs = generate_mcqs_from_summary_local(summary_text, num_qs=num_qs, model=MODEL)
        else:
            mcqs = []
        if produce_desc:
            descrs = generate_descriptive_with_answers("Video summary", context=summary_text, model=MODEL, num_qs=3)
        else:
            descrs = []

        results["Video summary"] = {"mcqs": mcqs, "descriptive": descrs}

        # Build files only containing the selected types
        df_all = build_dfs_from_questions(results)

        # CSV
        csv_bytes = df_all.to_csv(index=False).encode("utf-8")
        csv_key = hashlib.md5(csv_bytes).hexdigest()
        store_result_bytes(csv_key, csv_bytes, "video_questions.csv", "text/csv")

        # Excel
        excel_buf = BytesIO()
        with pd.ExcelWriter(excel_buf, engine="xlsxwriter") as writer:
            df_all.to_excel(writer, sheet_name="Questions", index=False)
        excel_buf.seek(0)
        excel_bytes = excel_buf.getvalue()
        excel_key = hashlib.md5(excel_bytes).hexdigest()
        store_result_bytes(excel_key, excel_bytes, "video_questions.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        # DOCX
        docx_bytes = build_docx_bytes(results)
        docx_key = hashlib.md5(docx_bytes).hexdigest()
        store_result_bytes(docx_key, docx_bytes, "video_questions.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # counts for this request
        mcq_count_now = len(mcqs)
        desc_count_now = len(descrs)

        # update global state
        STATE["mcq_count"] = STATE.get("mcq_count", 0) + mcq_count_now
        STATE["desc_count"] = STATE.get("desc_count", 0) + desc_count_now

        return {
            "status": "success",
            "mcqCount": mcq_count_now,
            "descCount": desc_count_now,
            "download_keys": {"csv": csv_key, "excel": excel_key, "docx": docx_key},
            "global_state": {
                "pdf_uploads": STATE["pdf_uploads"],
                "last_pdf_pages": STATE["last_pdf_pages"],
                "mcq_count": STATE["mcq_count"],
                "desc_count": STATE["desc_count"]
            },
            "results": results,
            "summary": summary_text,
            "chunks": chunk_summaries
        }
    except Exception as e:
        return {"status": "error", "error": str(e)}